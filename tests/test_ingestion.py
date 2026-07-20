"""文档摄取流水线测试。"""
from __future__ import annotations
import tempfile, time, unittest
from pathlib import Path
from unittest.mock import patch
from fastapi import FastAPI
from fastapi.testclient import TestClient

from service.api.documents import router
from service.ingestion.jobs import IngestionJobManager
from service.ingestion.registry import IngestionRegistry
from service.ingestion.service import IngestionError, IngestionService
from service.ingestion.parsers import assess_quality, parse_file
from service.ingestion.ocr import OcrUnavailable, call_ocr_parse


def _cfg(root: Path, knowledge: Path) -> dict:
    return {"paths": {"ingestion_dir": str(root)}, "ingestion": {"max_upload_bytes": 1024 * 1024,
        "allowed_extensions": [".md", ".txt", ".html", ".htm", ".docx", ".pdf"], "min_text_chars": 10, "auto_publish": False},
        "knowledge": {"domains": {"bank_stmt": {"enabled": True}}}}


class IngestionTests(unittest.TestCase):
    def test_html_docx_and_pdf_parsers(self):
        with tempfile.TemporaryDirectory() as tmp:
            directory = Path(tmp)
            html = directory / "sample.html"; html.write_text("<html><title>网页规则</title><h1>规则</h1><p>这是网页中的业务说明。</p></html>", encoding="utf-8")
            elements, _ = parse_file(html, html.name)
            self.assertTrue(any(e.text == "网页规则" for e in elements))

            from docx import Document
            doc = Document(); doc.add_heading("Word规则", level=1); doc.add_paragraph("这是Word中的业务说明。"); docx = directory / "sample.docx"; doc.save(docx)
            elements, meta = parse_file(docx, docx.name)
            self.assertGreater(meta["paragraphs"], 0); self.assertTrue(any("Word规则" in e.text for e in elements))

            from pypdf import PdfWriter
            writer = PdfWriter(); writer.add_blank_page(width=100, height=100); pdf = directory / "scan.pdf"
            with open(pdf, "wb") as f: writer.write(f)
            elements, meta = parse_file(pdf, pdf.name, {"enabled": False})
            quality = assess_quality(elements, ".pdf", 10)
            self.assertEqual(meta["page_count"], 1); self.assertIn("document_may_require_ocr", quality.warnings)

    def test_image_and_pdf_use_structured_company_ocr(self):
        with tempfile.TemporaryDirectory() as tmp:
            directory = Path(tmp)
            image_path = directory / "scan.png"; image_path.write_bytes(b"fake image")
            pages = [[
                {"type": "text", "text": "知识材料标题", "text_level": 1, "bbox": [1, 2, 3, 4], "page_idx": 1},
                {"type": "text", "text": "识别后的图片文字内容", "page_idx": 1},
                {"type": "table", "table_caption": "费用表", "table_body": "<table><tr><th>项目</th><th>金额</th></tr><tr><td>服务费</td><td>10</td></tr></table>", "page_idx": 1},
            ]]
            with patch("service.ingestion.ocr.call_ocr_parse", return_value=pages):
                elements, meta = parse_file(image_path, image_path.name, {"enabled": True, "flow": {"model": "br-ocr-v1"}})
            self.assertTrue(meta["ocr_used"]); self.assertEqual(meta["ocr_model"], "br-ocr-v1")
            self.assertEqual(elements[0].text, "知识材料标题")
            self.assertTrue(any("识别后" in e.text for e in elements))
            table = next(e for e in elements if e.type == "table")
            self.assertIn("| 项目 | 金额 |", table.text)
            self.assertEqual(table.metadata["ocr_raw"]["table_caption"], "费用表")

            from pypdf import PdfWriter
            writer = PdfWriter(); writer.add_blank_page(width=100, height=100); pdf = directory / "scan.pdf"
            with open(pdf, "wb") as f: writer.write(f)
            with patch("service.ingestion.ocr.call_ocr_parse", return_value=[[{"type": "text", "text": "扫描PDF识别文字内容", "page_idx": 1}]]):
                elements, meta = parse_file(pdf, pdf.name, {"enabled": True, "flow": {"model": "br-ocr-v1"}})
            self.assertTrue(meta["ocr_used"]); self.assertTrue(any("扫描PDF" in e.text for e in elements))

    def test_company_ocr_uploads_and_polls_due_agent_protocol(self):
        class FakeResponse:
            def __init__(self, payload): self.payload = payload
            def raise_for_status(self): return None
            def json(self): return self.payload

        calls = {"post": [], "get": []}
        poll_results = [
            FakeResponse({"data": {"task_status": "RUNNING"}}),
            FakeResponse({"data": {"task_status": "SUCCESS", "page_parse_json": "[{\"type\":\"text\",\"text\":\"结果\",\"page_idx\":2}]"}}),
        ]

        class FakeClient:
            def __init__(self, **kwargs): self.kwargs = kwargs
            def __enter__(self): return self
            def __exit__(self, *args): return False
            def post(self, url, json): calls["post"].append((url, json)); return FakeResponse({"code": 200, "data": {"task_uuid": "server-task-uuid"}})
            def get(self, url, **kwargs): calls["get"].append(url); return poll_results.pop(0)

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "材料.pdf"; path.write_bytes(b"pdf bytes")
            cfg = {"flow": {"base_url": "http://ocr.test", "model": "br-ocr-v1", "max_polls": 3, "poll_interval_seconds": 0}}
            with patch("service.ingestion.ocr.httpx.Client", FakeClient):
                pages = call_ocr_parse(path, cfg)

        self.assertEqual(pages[0][0]["text"], "结果")
        upload_url, body = calls["post"][0]
        self.assertEqual(upload_url, "http://ocr.test/api/v1/file_parse")
        self.assertEqual(body["model_name"], "br-ocr-v1")
        self.assertRegex(body["task_id"], r"^OCR\d{19}$")
        self.assertEqual(body["file_source_type"], "base64")
        self.assertEqual(body["split_setting"], {"split_method": "smart", "chunk_size": 300})
        self.assertEqual(len(calls["get"]), 2)
        self.assertTrue(all(url.endswith("/api/v1/result/server-task-uuid") for url in calls["get"]))

    def test_company_ocr_terminal_failure_is_explicit(self):
        class FakeResponse:
            def raise_for_status(self): return None
            def json(self): return {"data": {"task_status": "FAILED", "detail_msg": "模型解析失败"}}
        class FakeClient:
            def __init__(self, **kwargs): pass
            def __enter__(self): return self
            def __exit__(self, *args): return False
            def post(self, url, json): return FakeResponse()
            def get(self, url, **kwargs): return FakeResponse()
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "bad.pdf"; path.write_bytes(b"bad")
            with patch("service.ingestion.ocr.httpx.Client", FakeClient):
                with self.assertRaisesRegex(OcrUnavailable, "模型解析失败"):
                    call_ocr_parse(path, {"flow": {"max_polls": 1, "poll_interval_seconds": 0}})

    def test_pending_job_is_recovered(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "ingestion"; knowledge = Path(tmp) / "knowledge"; knowledge.mkdir()
            with patch("service.ingestion.service.get_config", return_value=_cfg(root, knowledge)):
                service = IngestionService(IngestionRegistry(str(root)))
            record = service.accept_upload("recover.md", "# 恢复\n\n这是一段可以恢复处理的知识内容。".encode(), "bank_stmt", {})
            service.registry.put_job("ingest_recover", {"job_id": "ingest_recover", "document_id": record["document_id"], "status": "running", "stage": "extracting", "error": None, "warnings": []})
            manager = IngestionJobManager(service); deadline = time.monotonic() + 2
            job = manager.get("ingest_recover")
            while job and job["status"] not in {"ready", "awaiting_review", "failed"} and time.monotonic() < deadline:
                time.sleep(0.01); job = manager.get("ingest_recover")
            self.assertEqual(job["status"], "ready")

    def test_markdown_lifecycle_and_deduplication(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "ingestion"; knowledge = Path(tmp) / "knowledge"; knowledge.mkdir()
            with patch("service.ingestion.service.get_config", return_value=_cfg(root, knowledge)):
                service = IngestionService(IngestionRegistry(str(root)))
            content = "# 余额规则\n\n## 定义\n\n余额必须与交易金额保持连续。".encode()
            record = service.accept_upload("余额.md", content, "bank_stmt", {"category": "交易错误"})
            canonical = service.process(record["document_id"])
            self.assertTrue(canonical.quality.passed)
            self.assertIn("# 余额规则", canonical.normalized_markdown)
            duplicate = service.accept_upload("另一个名字.md", content, "bank_stmt", {})
            self.assertTrue(duplicate["duplicate"])

            managed = Path(tmp) / "managed" / "bank_stmt"
            with patch("service.ingestion.service.get_managed_knowledge_dir", return_value=str(managed)):
                published = service.publish(record["document_id"])
                self.assertTrue(published.exists())
                self.assertEqual(published.parent, managed.resolve())
                service.delete(record["document_id"])
                self.assertFalse(published.exists())
            self.assertTrue(Path(record["raw_path"]).exists())

    def test_quality_gate_blocks_short_document(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "ingestion"; knowledge = Path(tmp) / "knowledge"; knowledge.mkdir()
            with patch("service.ingestion.service.get_config", return_value=_cfg(root, knowledge)):
                service = IngestionService(IngestionRegistry(str(root)))
            record = service.accept_upload("short.txt", b"tiny", "bank_stmt", {})
            service.process(record["document_id"])
            with patch("service.ingestion.service.get_managed_knowledge_dir", return_value=str(Path(tmp) / "managed" / "bank_stmt")):
                with self.assertRaises(IngestionError): service.publish(record["document_id"])

    def test_upload_preview_and_publish_api(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "ingestion"; knowledge = Path(tmp) / "knowledge"; knowledge.mkdir()
            with patch("service.ingestion.service.get_config", return_value=_cfg(root, knowledge)):
                service = IngestionService(IngestionRegistry(str(root)))
            manager = IngestionJobManager(service); app = FastAPI(); app.include_router(router)
            with patch("service.api.documents.get_ingestion_job_manager", return_value=manager), patch("service.ingestion.service.get_managed_knowledge_dir", return_value=str(Path(tmp) / "managed" / "bank_stmt")):
                with TestClient(app) as client:
                    response = client.post("/admin/documents", files={"file": ("rule.md", "# 规则\n\n## 内容\n\n这里是一段足够长的知识规则。", "text/markdown")}, data={"domain": "bank_stmt", "metadata": "{}"})
                    self.assertEqual(response.status_code, 202); payload = response.json()
                    deadline = time.monotonic() + 2
                    while time.monotonic() < deadline:
                        status = client.get(payload["status_url"]).json()
                        if status["status"] in {"ready", "awaiting_review", "failed"}: break
                        time.sleep(0.01)
                    self.assertEqual(status["status"], "ready")
                    preview = client.get(f"/admin/documents/{payload['document_id']}/preview")
                    self.assertEqual(preview.status_code, 200)
                    published = client.post(f"/admin/documents/{payload['document_id']}/publish?reindex=false")
                    self.assertEqual(published.status_code, 202)


if __name__ == "__main__": unittest.main()
