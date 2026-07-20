"""确定性文档解析器：输出统一 Element，不使用 LLM 改写事实。"""
from __future__ import annotations

import hashlib, re
from pathlib import Path
from typing import Any, Dict, List

from service.schemas.ingestion import DocumentElement, QualityReport


def _eid(index: int, text: str) -> str:
    return hashlib.sha256(f"{index}\0{text}".encode("utf-8")).hexdigest()[:24]


def _text_elements(text: str, title: str) -> List[DocumentElement]:
    elements: List[DocumentElement] = [DocumentElement(element_id=_eid(0, title), type="title", text=title, heading_path=[title])]
    headings = [title]
    buffer: List[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            idx = len(elements); kind = "list" if re.match(r"^\s*(?:[-*+] |\d+[.)] )", body) else "paragraph"
            elements.append(DocumentElement(element_id=_eid(idx, body), type=kind, text=body, heading_path=list(headings)))
        buffer.clear()

    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if match:
            flush(); level = len(match.group(1)); heading = match.group(2).strip()
            headings = headings[:level - 1] + [heading]
            idx = len(elements)
            elements.append(DocumentElement(element_id=_eid(idx, heading), type="heading", text=heading, heading_path=list(headings), metadata={"level": level}))
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
    flush()
    return elements


def _table_html_to_markdown(html: str) -> str:
    """将 OCR 的 table_body HTML 转成适合分块和向量化的 Markdown 表格。"""
    if not html.strip():
        return ""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    rows: List[List[str]] = []
    for row in soup.find_all("tr"):
        cells = [cell.get_text(" ", strip=True).replace("|", "\\|") for cell in row.find_all(["th", "td"])]
        if cells:
            rows.append(cells)
    if not rows:
        return soup.get_text(" ", strip=True)

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(normalized[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def _ocr_item_text(item: Dict[str, Any]) -> str:
    for key in ("text", "block_content", "content"):
        value = item.get(key)
        if value and str(value).strip():
            return str(value).strip()
    return ""


def _ocr_elements(pages: List[List[Dict[str, Any]]], fallback_title: str, model: str) -> tuple[List[DocumentElement], Dict[str, Any]]:
    title = fallback_title
    for page in pages:
        for item in page:
            text = _ocr_item_text(item)
            try:
                level = int(item.get("text_level", 0))
            except (TypeError, ValueError):
                level = 0
            if text and level == 1:
                title = text
                break
        if title != fallback_title:
            break

    elements = [DocumentElement(element_id=_eid(0, title), type="title", text=title, heading_path=[title])]
    headings: Dict[int, str] = {}
    confidences: List[float] = []
    page_numbers: List[int] = []

    for page_position, items in enumerate(pages, start=1):
        for item in items:
            try:
                page_no = max(1, int(item.get("page_idx", page_position)))
            except (TypeError, ValueError):
                page_no = page_position
            page_numbers.append(page_no)
            raw_type = str(item.get("type", "text")).lower()
            metadata: Dict[str, Any] = {"ocr_type": raw_type, "ocr_raw": item}
            score = item.get("confidence", item.get("score"))
            try:
                if score is not None:
                    confidences.append(float(score))
            except (TypeError, ValueError):
                pass

            if raw_type == "table":
                caption = str(item.get("table_caption") or item.get("table_title") or "").strip()
                footnote = str(item.get("table_footnote") or item.get("table_foot") or "").strip()
                table = _table_html_to_markdown(str(item.get("table_body") or ""))
                text = "\n\n".join(part for part in (caption, table, footnote) if part)
                element_type = "table"
            elif raw_type == "image":
                caption = str(item.get("image_caption") or item.get("image_title") or "").strip()
                footnote = str(item.get("image_footnote") or item.get("image_foot") or "").strip()
                text = "\n\n".join(part for part in (caption, footnote) if part)
                element_type = "paragraph"
            else:
                text = _ocr_item_text(item)
                element_type = "list" if re.match(r"^\s*(?:[-*+] |\d+[.)] )", text) else "paragraph"

            if not text:
                continue

            try:
                text_level = int(item.get("text_level", 0))
            except (TypeError, ValueError):
                text_level = 0
            if raw_type == "text" and text_level > 0:
                if text_level == 1 and text == title:
                    continue
                level = max(2, min(6, text_level))
                headings[level] = text
                headings = {key: value for key, value in headings.items() if key <= level}
                heading_path = [title] + [headings[key] for key in sorted(headings)]
                metadata["level"] = level
                element_type = "heading"
            else:
                heading_path = [title] + [headings[key] for key in sorted(headings)]

            index = len(elements)
            elements.append(DocumentElement(
                element_id=_eid(index, f"{page_no}\0{text}"),
                type=element_type,
                text=text,
                page=page_no,
                heading_path=heading_path,
                metadata=metadata,
            ))

    page_count = max(page_numbers, default=len(pages))
    return elements, {
        "page_count": page_count,
        "ocr_used": True,
        "ocr_provider": "company_flow",
        "ocr_model": model,
        "ocr_confidence": round(sum(confidences) / len(confidences), 2) if confidences else None,
    }


def parse_file(path: Path, source_filename: str, ocr_config: Dict[str, Any] | None = None) -> tuple[List[DocumentElement], Dict[str, Any]]:
    ext = path.suffix.lower(); title = Path(source_filename).stem
    if ext in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8-sig")
        if ext == ".md":
            from service.knowledge.markdown_parser import parse_markdown
            parsed = parse_markdown(text); text = parsed.content
            first_h1 = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
            title = str(parsed.frontmatter.get("title") or (first_h1.group(1).strip() if first_h1 else title))
            if first_h1 and first_h1.group(1).strip() == title:
                text = text[:first_h1.start()] + text[first_h1.end():]
        return _text_elements(text, title), {}
    if ext in {".html", ".htm"}:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(path.read_text(encoding="utf-8-sig"), "html.parser")
        for node in soup(["script", "style", "nav"]): node.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else title
        lines = []
        for node in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
            value = node.get_text(" ", strip=True)
            if not value: continue
            if node.name.startswith("h"): lines.append(f"{'#' * int(node.name[1])} {value}")
            elif node.name == "li": lines.append(f"- {value}")
            else: lines.append(value)
        return _text_elements("\n\n".join(lines), title), {}
    if ext == ".docx":
        from docx import Document
        doc = Document(str(path)); lines: List[str] = []
        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if not value: continue
            style = paragraph.style.name.lower() if paragraph.style else ""
            match = re.search(r"heading\s*(\d+)", style)
            lines.append(f"{'#' * min(6, int(match.group(1)))} {value}" if match else value)
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                lines.append("\n".join("| " + " | ".join(row) + " |" for row in rows))
        return _text_elements("\n\n".join(lines), title), {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
    if ext == ".pdf":
        import fitz
        cfg = ocr_config or {}
        if cfg.get("enabled", True):
            from service.ingestion.ocr import call_ocr_parse
            pages = call_ocr_parse(path, cfg)
            model = str(cfg.get("flow", cfg).get("model", "br-ocr-v1"))
            return _ocr_elements(pages, title, model)

        elements = [DocumentElement(element_id=_eid(0, title), type="title", text=title, heading_path=[title])]
        with fitz.open(path) as pdf:
            page_count = len(pdf)
            for page_no, page in enumerate(pdf, start=1):
                text = page.get_text("text").strip()
                if text:
                    elements.append(DocumentElement(element_id=_eid(len(elements), text), type="paragraph", text=text, page=page_no, heading_path=[title], metadata={"ocr": False}))
        return elements, {"page_count": page_count, "ocr_used": False}
    if ext in {".png", ".jpg", ".jpeg"}:
        cfg = ocr_config or {}
        if not cfg.get("enabled", True):
            return _text_elements("", title), {"ocr_used": False, "ocr_error": "OCR disabled"}
        from service.ingestion.ocr import call_ocr_parse
        pages = call_ocr_parse(path, cfg)
        model = str(cfg.get("flow", cfg).get("model", "br-ocr-v1"))
        return _ocr_elements(pages, title, model)
    raise ValueError(f"unsupported file extension: {ext}")


def render_markdown(elements: List[DocumentElement], fallback_title: str) -> str:
    parts: List[str] = []
    has_h1 = False
    for element in elements:
        if element.type == "title":
            if not has_h1: parts.append(f"# {element.text}"); has_h1 = True
        elif element.type == "heading":
            level = int(element.metadata.get("level", 2)); parts.append(f"{'#' * max(2, min(6, level))} {element.text}")
        else: parts.append(element.text)
    return "\n\n".join(parts or [f"# {fallback_title}"]).strip() + "\n"


def assess_quality(elements: List[DocumentElement], source_type: str, min_chars: int, parser_metadata: Dict[str, Any] | None = None, min_ocr_confidence: float = 60) -> QualityReport:
    chars = sum(len(e.text.strip()) for e in elements if e.type not in {"title", "heading"})
    warnings: List[str] = []
    if chars < min_chars: warnings.append("extracted_text_too_short")
    meta = parser_metadata or {}
    if source_type in {".pdf", ".png", ".jpg", ".jpeg"} and chars < max(200, min_chars): warnings.append("document_may_require_ocr")
    if meta.get("ocr_error"): warnings.append("ocr_unavailable")
    if meta.get("ocr_confidence") is not None and float(meta["ocr_confidence"]) < min_ocr_confidence: warnings.append("low_ocr_confidence")
    if chars > 2000 and not any(e.type == "heading" for e in elements): warnings.append("long_document_without_headings")
    requires_review = bool(warnings)
    return QualityReport(passed=chars >= min_chars, requires_review=requires_review, text_chars=chars, element_count=len(elements), warnings=warnings)
